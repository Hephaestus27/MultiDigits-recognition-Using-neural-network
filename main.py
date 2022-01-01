#

# Author : Hephaestus27 : https://github.com/Hephaestus27

# #TF compiling problems :
import os


os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'

from tensorflow.keras.models import load_model
import cv2
import numpy as np
import matplotlib.pyplot as plt


model = load_model('handwritten.model')

line = 0
def IsSameLine(Y_final,lineindicator): #XO : Old X
    if( abs(lineindicator*1.0 - Y_final) > (60) ):
        global line
        line = Y_final #new line indicator
        return 0
    else:
        return 1






path = './raw_images/1.jpg'
basename = os.path.basename(path)  #get the file name and ext
basename = basename.split('.')[0]
image = cv2.imread(path)


# Step 1 : Converting to grayscale
grey = cv2.cvtColor(image.copy(), cv2.COLOR_BGR2GRAY)
# Step 2 : Binarize(threshold) the greyscaled image in such a way that only the digits in the image are white and rest is black
ret, thresh = cv2.threshold(grey.copy(), 75, 255, cv2.THRESH_BINARY_INV)
# Step 3 :  Using the binarized image, find contours in the image. Here, contours will provide us the individual digits in the image
contours, __ = cv2.findContours(thresh.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
preprocessed_digits = []
labeled_preprocessed_digits = []
line=0
count = 0
for c in contours: #Detect Multiple Digits
    x,y,w,h = cv2.boundingRect(c)

    if(count == 1):
        line = y+h # get the line indicator
    count +=1
    # Creating a rectangle around the digit in the original image (for displaying the digits fetched via contours)
    cv2.rectangle(image, (x,y), (x+w, y+h), color=(0, 255, 0), thickness=2)

#debug :
    #print("X_origin : ",x)
    #print("X_Final : ",x+w)
    #print(" W :",w)
    #print("Y_origin : ",y)
    #print("Y_Final : ",y+h)
    #print(" H :",h)
    #print('\n')


# Cropping out the digit from the image corresponding to the current contours in the for loop
    digit = thresh[y:y+h, x:x+w]
    # Resizing that digit to (18, 18)
    resized_digit = cv2.resize(digit, (18,18))
    # Padding the digit with 5 pixels of black color (zeros) in each side to finally produce the image of (28, 28)
    padded_digit = np.pad(resized_digit, ((5,5),(5,5)), "constant", constant_values=0)
    # Adding the preprocessed digit to the list of preprocessed digits
    preprocessed_digits.append(padded_digit)
    final = y+h
    labeled_preprocessed_digits.append(IsSameLine(final,line))
plt.imshow(image, cmap="gray")
plt.show()

inp = np.array(preprocessed_digits)


#Prediction
#arr = np.array([])  #append result in array
Returned_String =''
count = 0
for digit in preprocessed_digits:
    prediction = model.predict(digit.reshape(1, 28, 28, 1))
    #plt.imshow(digit.reshape(28, 28), cmap="gray")
    #plt.show()
    # arr = np.append(arr,format(np.argmax(prediction)))
    if(labeled_preprocessed_digits[count] == 0):
        Returned_String += '%'
    Returned_String += (str(np.argmax(prediction)))
    count+=1

Final_Doc_Data = Returned_String.split('%') # i decided to add % when ever a new line is here
#save on a new txt file

#output = basename + '.txt'
#f = open(output, "w")
#f.write(Returned_String)

#to append to a word file
from docx import Document
from docx.enum.text import WD_BREAK

DocName = basename +'.docx'
document = Document()
document.add_heading('OCR for the file : ' + basename, 0)
for DocumentLine in Final_Doc_Data :
    doc_para = document.add_paragraph(DocumentLine)

document.save(DocName)
print('" '+DocName +'" has been saved')